import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell internal margins/padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    """Set subtle borders for table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def add_styled_heading(doc, text, level):
    """Add beautifully styled headings with colors and spacing."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(27, 54, 93) # Navy
        # Add subtle bottom accent rule under Heading 1
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="2563EB"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run.font.size = Pt(12.5)
        run.font.color.rgb = RGBColor(43, 76, 126) # Slate Blue
    elif level == 3:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate Dark
    return p

def add_body_paragraph(doc, text="", bold_prefix="", italic=False):
    """Add clean body paragraph with 1.15 line spacing and 4pt after spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = "Arial"
        r_prefix.font.size = Pt(10)
        r_prefix.bold = True
        r_prefix.font.color.rgb = RGBColor(15, 23, 42)
        
    if text:
        r_text = p.add_run(text)
        r_text.font.name = "Arial"
        r_text.font.size = Pt(10)
        r_text.italic = italic
        r_text.font.color.rgb = RGBColor(51, 65, 85)
        
    return p

def add_bullet_point(doc, text, bold_title=""):
    """Add a beautifully indented bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_title:
        r_title = p.add_run(bold_title)
        r_title.font.name = "Arial"
        r_title.font.size = Pt(10)
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(15, 23, 42)
        
    r_text = p.add_run(text)
    r_text.font.name = "Arial"
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, title, text):
    """Add an elegant callout box with a colored left border."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F7FF")
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    # Left thick border
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="2563EB"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"KEY TAKEAWAY: {title}\n")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(9.5)
    r_title.bold = True
    r_title.font.color.rgb = RGBColor(30, 64, 175)
    
    r_text = p.add_run(text)
    r_text.font.name = "Arial"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(30, 41, 59)
    
    # Spacing after table
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_str, caption=""):
    """Add a shaded monospace code snippet box with caption."""
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(8)
        p_cap.paragraph_format.space_after = Pt(3)
        p_cap.paragraph_format.keep_with_next = True
        r_cap = p_cap.add_run(f"Listing: {caption}")
        r_cap.font.name = "Arial"
        r_cap.font.size = Pt(9.5)
        r_cap.bold = True
        r_cap.font.color.rgb = RGBColor(43, 76, 126)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/><w:top w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/><w:right w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/><w:bottom w:val="single" w:sz="8" w:space="0" w:color="CBD5E1"/></w:tcBorders>')
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    
    r_code = p.add_run(code_str.strip())
    r_code.font.name = "Consolas"
    r_code.font.size = Pt(8.5)
    r_code.font.color.rgb = RGBColor(15, 23, 42)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def add_styled_table(doc, headers, data, col_widths=None, alignment=None):
    """Add a professional table with dark header, subtle borders, and zebra rows."""
    tbl = doc.add_table(rows=len(data) + 1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    set_table_borders(tbl, color="CBD5E1", sz="4")
    
    # Format Header Row
    hdr_cells = tbl.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3E62") # Dark Blue
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if alignment and len(alignment) > i:
            p.alignment = alignment[i]
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Format Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = tbl.rows[r_idx + 1].cells
        row_bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], row_bg)
            set_cell_margins(row_cells[c_idx], top=90, bottom=90, left=120, right=120)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            if alignment and len(alignment) > c_idx:
                p.alignment = alignment[c_idx]
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(30, 41, 59)
                
    # Apply column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
                
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(0)
    p_spacer.paragraph_format.space_after = Pt(6)

def add_image_with_caption(doc, image_path, caption_text, width_inches=6.2):
    """Add a centered image with a styled caption."""
    if not os.path.exists(image_path):
        p_err = doc.add_paragraph(f"[Image placeholder: {caption_text}]")
        p_err.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    p_img = doc.add_paragraph()
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(4)
    p_img.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.keep_with_next = True
    
    run_img = p_img.add_run()
    run_img.add_picture(image_path, width=Inches(width_inches))
    
    p_cap = doc.add_paragraph()
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(10)
    p_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_cap = p_cap.add_run(caption_text)
    r_cap.font.name = "Arial"
    r_cap.font.size = Pt(9)
    r_cap.font.italic = True
    r_cap.font.bold = True
    r_cap.font.color.rgb = RGBColor(71, 85, 105)

print("Document helper utilities loaded successfully.")
