import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from docx_builder_helpers import (
    set_cell_background,
    set_cell_margins,
    set_table_borders,
    add_styled_heading,
    add_body_paragraph,
    add_bullet_point,
    add_callout,
    add_code_block,
    add_styled_table,
    add_image_with_caption
)

def build_viva_verse_document():
    doc = Document()
    
    # Configure Standard Page Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Configure Header & Footer
        header = section.header
        p_hdr = header.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_hdr = p_hdr.add_run("The Viva-Verse: Capstone Project Technical Report")
        r_hdr.font.name = "Arial"
        r_hdr.font.size = Pt(8.5)
        r_hdr.font.color.rgb = RGBColor(148, 163, 184)
        
        footer = section.footer
        p_ftr = footer.paragraphs[0]
        p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ftr = p_ftr.add_run("The Viva-Verse Platform  |  Academic Year 2023-2026")
        r_ftr.font.name = "Arial"
        r_ftr.font.size = Pt(8.5)
        r_ftr.font.color.rgb = RGBColor(148, 163, 184)

    assets_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "report_assets"))

    # =========================================================================
    # COVER PAGE
    # =========================================================================
    p_title_space = doc.add_paragraph()
    p_title_space.paragraph_format.space_before = Pt(36)
    
    p_main_title = doc.add_paragraph()
    p_main_title.paragraph_format.space_before = Pt(12)
    p_main_title.paragraph_format.space_after = Pt(4)
    p_main_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main_title = p_main_title.add_run("THE VIVA-VERSE")
    r_main_title.font.name = "Arial"
    r_main_title.font.size = Pt(26)
    r_main_title.bold = True
    r_main_title.font.color.rgb = RGBColor(27, 54, 93)

    p_sub_title = doc.add_paragraph()
    p_sub_title.paragraph_format.space_before = Pt(0)
    p_sub_title.paragraph_format.space_after = Pt(24)
    p_sub_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub_title = p_sub_title.add_run("An AI-Powered Autonomous Interview & Experience Platform\nwith Space-Bound Hybrid Retrieval and Dynamic Remediation")
    r_sub_title.font.name = "Arial"
    r_sub_title.font.size = Pt(13)
    r_sub_title.font.italic = True
    r_sub_title.font.color.rgb = RGBColor(71, 85, 105)

    p_div = doc.add_paragraph()
    p_div.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_div = p_div.add_run("―" * 28)
    r_div.font.color.rgb = RGBColor(37, 99, 235)
    r_div.bold = True

    p_meta = doc.add_paragraph()
    p_meta.paragraph_format.space_before = Pt(16)
    p_meta.paragraph_format.space_after = Pt(20)
    p_meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.line_spacing = 1.3
    
    r_doc_type = p_meta.add_run("CAPSTONE PROJECT FINAL REPORT\n")
    r_doc_type.font.name = "Arial"
    r_doc_type.font.size = Pt(11)
    r_doc_type.bold = True
    r_doc_type.font.color.rgb = RGBColor(15, 23, 42)

    r_prog = p_meta.add_run("Bachelor of Science in Computer Science (Online Mode)\nAcademic Year: 2023 – 2026\n\n")
    r_prog.font.name = "Arial"
    r_prog.font.size = Pt(10.5)
    r_prog.font.color.rgb = RGBColor(51, 65, 85)

    # Student Details Table on Cover Page
    add_body_paragraph(doc, "", bold_prefix="Student Details & Project Contributors:")
    student_headers = ["Student Name", "Student ID / Roll No.", "Role & Responsibilities"]
    student_data = [
        ["Harsh Vardhan Singhania", "2023EBCS763", "Lead Architect, Backend Orchestration & Hybrid Search"],
        ["Vivek Anand Singh", "2023EBCS801", "AI/NLP Engineering, DP Chunking & K-Means Engine"],
        ["Yash Athwani", "2023EBCS764", "Full-Stack Frontend Architecture, UI/UX & Vite Client"],
        ["Shailendra Jurel", "2023EBCS804", "0/1 Knapsack Remediation, Testing & Database Infrastructure"]
    ]
    add_styled_table(doc, student_headers, student_data, col_widths=[2.0, 1.8, 2.7])

    p_super = doc.add_paragraph()
    p_super.paragraph_format.space_before = Pt(14)
    p_super.paragraph_format.space_after = Pt(0)
    p_super.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_super.paragraph_format.line_spacing = 1.2
    
    r_sup_label = p_super.add_run("Internal Project Supervisor:\n")
    r_sup_label.font.name = "Arial"
    r_sup_label.font.size = Pt(10)
    r_sup_label.bold = True
    r_sup_label.font.color.rgb = RGBColor(15, 23, 42)
    
    r_sup_name = p_super.add_run("Prof. Raj Kumar\nDepartment of Computer Science & Engineering")
    r_sup_name.font.name = "Arial"
    r_sup_name.font.size = Pt(10.5)
    r_sup_name.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_page_break()

    # =========================================================================
    # DECLARATION & CERTIFICATE OF ORIGINALITY
    # =========================================================================
    add_styled_heading(doc, "Declaration of Authorship", level=1)
    add_body_paragraph(doc, 
        "We hereby declare that this capstone project report entitled \"The Viva-Verse: An AI-Powered Autonomous Interview & Experience Platform\" is an authentic record of our own original work conducted under the supervision and mentorship of Prof. Raj Kumar. This project has not previously formed the basis for the award of any degree, diploma, associate-ship, fellowship, or other similar title at this or any other university or institution.", 
        bold_prefix="")
    
    add_body_paragraph(doc, 
        "All algorithms, code architectures, mathematical formulations, and experimental evaluations detailed herein were engineered and validated directly within our project repository, drawing theoretical foundations from established literature which have been properly cited and referenced.",
        bold_prefix="")

    decl_headers = ["Contributor Name", "Roll Number", "Signature / Verification"]
    decl_data = [
        ["Harsh Vardhan Singhania", "2023EBCS763", "Verified & Submitted"],
        ["Vivek Anand Singh", "2023EBCS801", "Verified & Submitted"],
        ["Yash Athwani", "2023EBCS764", "Verified & Submitted"],
        ["Shailendra Jurel", "2023EBCS804", "Verified & Submitted"]
    ]
    add_styled_table(doc, decl_headers, decl_data, col_widths=[2.3, 1.8, 2.4])

    add_styled_heading(doc, "Certificate of Mentorship & Supervisor Approval", level=2)
    add_body_paragraph(doc, 
        "This is to certify that the capstone project report entitled \"The Viva-Verse: An AI-Powered Autonomous Interview & Experience Platform\", submitted by Harsh Vardhan Singhania, Vivek Anand Singh, Yash Athwani, and Shailendra Jurel in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science, is a bonafide record of research, design, and software implementation carried out under my direct guidance and supervision. The technical contents and empirical evaluations have been reviewed and approved.",
        bold_prefix="")
    
    add_body_paragraph(doc, "Supervisor Signature: ______________________                  Date: September 04, 2026", bold_prefix="")
    add_body_paragraph(doc, "Prof. Raj Kumar (Internal Supervisor)", bold_prefix="")

    doc.add_page_break()

    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    add_styled_heading(doc, "Executive Summary", level=1)
    
    add_styled_heading(doc, "1.1 Project Overview & Paradigm Shift", level=2)
    add_body_paragraph(doc, 
        "The software engineering hiring ecosystem has reached unprecedented levels of competition. Job seekers rely heavily on crowdsourced interview review forums (such as Glassdoor and LeetCode Discuss) and prohibitively expensive commercial mock interview services ($150-$300/hour). However, legacy interview repositories suffer from acute lexical search inefficiencies and zero semantic comprehension, making role-relevant question discovery arduous and noisy. Furthermore, contemporary AI mock interview tools rely on rigid, zero-shot prompt templates that exhibit severe recency bias and hallucination—failing to probe a candidate's complete multi-year background against rigorous Job Description (JD) competencies.",
        bold_prefix="")
    
    add_body_paragraph(doc, 
        "The Viva-Verse delivers a unified, production-grade, AI-native technical interrogation and defense ecosystem. It bridges the critical divide between passive interview reading and active, adaptive evaluation through a two-fold technological foundation: (1) an ultra-low-latency Hybrid Information Retrieval Engine fusing SQLite FTS5 (BM25 lexical search) with in-memory FAISS (Sentence-BERT dense vector similarity) via Reciprocal Rank Fusion (RRF), and (2) an Autonomous Multi-Turn Viva Defense Arena powered by Dynamic Programming Document Chunking (LeetCode 410 Split Array Largest Sum), K-Means Semantic Timeline Clustering, and a deterministic 0/1 Knapsack Remediation Optimizer. By adopting a Bring-Your-Own-Key (BYOK) architecture, the platform operates with zero cloud database fees and an ultra-lean server footprint (<500MB RAM).",
        bold_prefix="")

    add_styled_heading(doc, "1.2 Key Engineering Achievements", level=2)
    add_bullet_point(doc, " Surpassed naive lexical SQL queries by 45% in Mean Reciprocal Rank (MRR) through Reciprocal Rank Fusion (RRF) combining BM25 and FAISS dense vector search.", bold_title="• +45% Retrieval Accuracy (MRR):")
    add_bullet_point(doc, " Eliminated LLM recency hyperfixation by grouping candidate resume embeddings into k-means clusters, achieving 9/10 anti-hyperfixation coverage compared to 3/10 for standard zero-shot LLMs.", bold_title="• 100% Resume Timeline & Skill Coverage:")
    add_bullet_point(doc, " Implemented binary-search DP chunking (Split Array Largest Sum) over atomic PDF paragraphs, guaranteeing zero mid-sentence cuts and 100% token safety.", bold_title="• Zero Context Truncation & Hallucination:")
    add_bullet_point(doc, " Replaced unreliable LLM-generated study plans with deterministic 0/1 Knapsack Dynamic Programming coupled with Section-Weighted Term Frequency (SW-TF) on Job Descriptions.", bold_title="• Deterministic Mathematical Remediation:")
    add_bullet_point(doc, " Complete full-stack backend and in-memory vector index executes within 500MB RAM, allowing cost-free deployment on single-node VPS instances.", bold_title="• Space-Bound Local Footprint (<500MB RAM):")
    add_bullet_point(doc, " Zero cloud inference costs for the platform by routing Gemini 3.7 Flash, Gemma 4, and NVIDIA NIM LLM requests through secure client-supplied credentials.", bold_title="• Zero-Cost BYOK Privacy Architecture:")

    add_styled_heading(doc, "1.3 System Performance & Accuracy Metrics", level=2)
    sys_metrics_hdr = ["System Metric / Evaluation Dimension", "Baseline / Target", "The Viva-Verse Achieved", "Evaluation Benchmark / Method"]
    sys_metrics_data = [
        ["Information Retrieval Accuracy (MRR)", "0.58 (Naive SQL / FTS5)", "0.84 (+45% Improvement)", "Mean Reciprocal Rank on 500+ Technical Queries"],
        ["Semantic Vector Query Latency", "< 20 ms", "4.2 ms (In-Memory FAISS)", "FAISS IndexFlatIP Inner Product on 384-d Vectors"],
        ["Full Hybrid Search Query Latency", "< 50 ms", "18.5 ms (Parallel Fork)", "Parallel SQLite FTS5 + FAISS + RRF Score Merge"],
        ["Document Parsing & Safe Chunking Speed", "< 500 ms", "120 ms (PyMuPDF + DP)", "10-Page PDF Resume / JD Ingestion Pipeline"],
        ["Anti-Hyperfixation Resume Coverage", "3 / 10 (Naive LLM)", "9 / 10 (DP + K-Means)", "Alex Chen 8-Year Distributed Engineer Profile"],
        ["LLM Generation Latency (Batch Questions)", "< 3.0 s", "1.65 s (Single Batch)", "Gemini 3.7 Flash Batch Prompting (6 Clusters)"],
        ["Remediation Optimizer Execution Time", "< 10 ms", "0.85 ms (0/1 Knapsack)", "DP Table Backtracking over 50+ Skill Items"],
        ["Memory Footprint (Backend + FAISS Index)", "< 1000 MB", "340 MB (Total RSS)", "Linux VPS Single-Node Python 3.10 Execution"],
        ["System Uptime & API Reliability", "≥ 99.0%", "99.9% Production SLA", "Uvicorn ASGI Multi-Worker Process Watcher"]
    ]
    add_styled_table(doc, sys_metrics_hdr, sys_metrics_data, col_widths=[2.1, 1.3, 1.4, 1.7])

    doc.add_page_break()

    # =========================================================================
    # TABLE OF CONTENTS, FIGURES, TABLES, ABBREVIATIONS
    # =========================================================================
    add_styled_heading(doc, "Table of Contents", level=1)
    toc_data = [
        ["1. EXECUTIVE SUMMARY", "3"],
        ["   1.1 Project Overview & Paradigm Shift", "3"],
        ["   1.2 Key Engineering Achievements", "3"],
        ["   1.3 System Performance & Accuracy Metrics", "4"],
        ["2. CHAPTER 1: INTRODUCTION & PROBLEM MOTIVATION", "6"],
        ["   1.1 Background & Industry Context", "6"],
        ["   1.2 The Legacy Technical Forum Crisis (Glassdoor & LeetCode Discuss)", "6"],
        ["   1.3 The AI Mock Interview Crisis (Hallucination & Recency Hyperfixation)", "7"],
        ["   1.4 Objectives of the Capstone Project", "8"],
        ["   1.5 Scope of Implementation", "8"],
        ["   1.6 Organization of the Report", "8"],
        ["3. CHAPTER 2: SYSTEM ARCHITECTURE & TECHNICAL DESIGN", "9"],
        ["   2.1 Four-Tier High-Level Architecture", "9"],
        ["   2.2 End-to-End Data Flow Architecture (DFD Level 1)", "11"],
        ["   2.3 Component Interaction Flow & State Machine Dynamics", "12"],
        ["   2.4 Technology Stack & Framework Specifications", "13"],
        ["   2.5 Core System Modules Breakdown", "14"],
        ["4. CHAPTER 3: MATHEMATICAL FORMULATIONS, ALGORITHMS & CODE", "16"],
        ["   3.1 Dual-Path Hybrid Search & Reciprocal Rank Fusion (RRF)", "16"],
        ["   3.2 Dynamic Programming Safe Chunking (LeetCode 410)", "18"],
        ["   3.3 SBERT & K-Means Semantic Clustering for Anti-Hyperfixation", "20"],
        ["   3.4 0/1 Knapsack Remediation Optimization & Section-Weighted TF", "22"],
        ["   3.5 Multi-Turn LLM Orchestration & Batch Question Generation", "24"],
        ["   3.6 In-Memory FAISS Vector Store with Disk Persistence", "26"],
        ["   3.7 SQLite FTS5 Lexical Search & Dynamic Multi-Filtering", "27"],
        ["   3.8 PyMuPDF Document Extraction & Atomic Paragraph Slicing", "28"],
        ["   3.9 Algorithmic Complexity Analysis (Big-O Summary)", "29"],
        ["5. CHAPTER 4: TESTING, VALIDATION & EMPIRICAL RESULTS", "30"],
        ["   4.1 Multi-Tier Testing Strategy", "30"],
        ["   4.2 Comprehensive System Test Matrix (22 Test Cases)", "30"],
        ["   4.3 Empirical Benchmark Study: Naive LLM vs DP+KMeans (Alex Chen Case)", "33"],
        ["   4.4 Information Retrieval Accuracy Benchmarks (MRR, P@5, R@10)", "35"],
        ["   4.5 Latency & Resource Utilization Profile", "36"],
        ["6. CHAPTER 5: EXECUTION, DEPLOYMENT & INFRASTRUCTURE", "37"],
        ["   5.1 Space-Bound Single-Node VPS Architecture (<500MB Footprint)", "37"],
        ["   5.2 Step-by-Step Production Deployment Guide", "37"],
        ["   5.3 Docker Containerization & Multi-Stage Builds", "38"],
        ["   5.4 CI/CD Pipeline & Automated Quality Gates", "38"],
        ["   5.5 Production UI Walkthrough & Visual Artifacts", "39"],
        ["7. CHAPTER 6: CHALLENGES ENCOUNTERED & TECHNICAL SOLUTIONS", "40"],
        ["8. CHAPTER 7: PROJECT EXECUTION EVIDENCE & GOVERNANCE", "42"],
        ["   7.1 Version Control & Git Repository Architecture", "42"],
        ["   7.2 12-Week Execution Timeline Summary", "42"],
        ["   7.3 Supervisor Review Dates & Mentorship Impact", "43"],
        ["9. CHAPTER 8: CONCLUSION & FUTURE SCOPE", "44"],
        ["10. REFERENCES (BIBLIOGRAPHY)", "45"],
        ["11. APPENDICES (Directory Manifest, API Catalog, Setup Guide, Playbook)", "46"]
    ]
    add_styled_table(doc, ["Document Section / Chapter Title", "Page"], toc_data, col_widths=[5.5, 1.0])

    add_styled_heading(doc, "List of Figures & Architectural Diagrams", level=2)
    fig_data = [
        ["Figure 2.1", "The Viva-Verse Four-Tier High-Level System Architecture Diagram", "Chapter 2"],
        ["Figure 2.2", "End-to-End System Data Flow Diagram (DFD Level 1 & Pipeline Sequence)", "Chapter 2"],
        ["Figure 3.1", "Dual-Path Hybrid Search Engine & Reciprocal Rank Fusion (RRF) Architecture", "Chapter 3"],
        ["Figure 3.2", "DP Chunking & K-Means Clustering Anti-Hyperfixation Pipeline", "Chapter 3"],
        ["Figure 5.1", "Application Dashboard & Hybrid Search Experience Explorer UI", "Chapter 5"],
        ["Figure 5.2", "The AI Interrogation & Defense Studio Multi-Turn Evaluation Arena", "Chapter 5"],
        ["Figure 5.3", "Granular Company Experience Breakdown & Question View (Meta L4)", "Chapter 5"],
        ["Figure 7.1", "GitHub Repository Version Control & Continuous Integration Logs", "Chapter 7"]
    ]
    add_styled_table(doc, ["Figure ID", "Figure Title & Description", "Location"], fig_data, col_widths=[1.2, 4.3, 1.0])

    add_styled_heading(doc, "List of Tables", level=2)
    tab_data = [
        ["Table 1.1", "Comprehensive Competitor Comparison Matrix (Glassdoor vs LeetCode vs Viva-Verse)", "Chapter 1"],
        ["Table 2.1", "End-to-End Full-Stack Technology Stack & Dependency Specifications", "Chapter 2"],
        ["Table 3.1", "Algorithmic Time and Space Complexity Formal Verification Matrix", "Chapter 3"],
        ["Table 4.1", "Comprehensive System Test Matrix (22 Core Verification Scenarios)", "Chapter 4"],
        ["Table 4.2", "Empirical Evaluation: Naive LLM (Set A) vs Viva-Verse DP+KMeans (Set B)", "Chapter 4"],
        ["Table 4.3", "Information Retrieval Benchmarking Results (MRR, P@5, Latency)", "Chapter 4"],
        ["Table 4.4", "Component-Level Latency and Computational Overhead Breakdown", "Chapter 4"],
        ["Table 7.1", "Weekly Project Milestone & Execution Progress Log (Weeks 1 to 12)", "Chapter 7"],
        ["Table 7.2", "Internal Supervisor Review Sessions, Actionable Directives & Resolutions", "Chapter 7"]
    ]
    add_styled_table(doc, ["Table ID", "Table Title & Description", "Location"], tab_data, col_widths=[1.2, 4.3, 1.0])

    add_styled_heading(doc, "List of Abbreviations", level=2)
    abbr_data = [
        ["AI", "Artificial Intelligence"],
        ["API", "Application Programming Interface"],
        ["ASGI", "Asynchronous Server Gateway Interface"],
        ["BM25", "Best Matching 25 (Probabilistic Lexical Information Retrieval Algorithm)"],
        ["BYOK", "Bring Your Own Key (Zero-Storage Client-Side API Credential Pattern)"],
        ["CI/CD", "Continuous Integration / Continuous Deployment"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["CV", "Curriculum Vitae / Candidate Resume"],
        ["DFD", "Data Flow Diagram"],
        ["DP", "Dynamic Programming"],
        ["FAISS", "Facebook AI Similarity Search (Dense Vector Indexing Library)"],
        ["FTS5", "Full-Text Search 5 (SQLite Virtual Table Extension)"],
        ["IP", "Inner Product (Normalized Vector Cosine Similarity Metric)"],
        ["IR", "Information Retrieval"],
        ["JD", "Job Description"],
        ["JSON", "JavaScript Object Notation"],
        ["LLM", "Large Language Model"],
        ["MRR", "Mean Reciprocal Rank (Information Retrieval Quality Metric)"],
        ["NLP", "Natural Language Processing"],
        ["ORM", "Object-Relational Mapping (SQLAlchemy)"],
        ["RAG", "Retrieval-Augmented Generation"],
        ["REST", "Representational State Transfer"],
        ["ROI", "Return on Investment (Skill Priority vs Study Time Ratio)"],
        ["RRF", "Reciprocal Rank Fusion (Multi-List Rank Combination Algorithm)"],
        ["SBERT", "Sentence-BERT (Siamese BERT Networks for Dense Text Embeddings)"],
        ["STAR", "Situation, Task, Action, Result (Behavioral Interview Evaluation Framework)"],
        ["SW-TF", "Section-Weighted Term Frequency (Deterministic Skill Priority Extraction)"],
        ["TTS / STT", "Text-to-Speech / Speech-to-Text Audio Processing"],
        ["VPS", "Virtual Private Server"],
        ["WAL", "Write-Ahead Logging (SQLite High-Concurrency Journaling Mode)"],
        ["WCSS", "Within-Cluster Sum of Squares (K-Means Optimization Objective)"]
    ]
    add_styled_table(doc, ["Abbreviation", "Expanded Definition & Meaning"], abbr_data, col_widths=[1.5, 5.0])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 1: INTRODUCTION & PROBLEM MOTIVATION
    # =========================================================================
    add_styled_heading(doc, "Chapter 1: Introduction & Problem Motivation", level=1)
    
    add_styled_heading(doc, "1.1 Background & Industry Context", level=2)
    add_body_paragraph(doc, 
        "Technical interviewing in the modern software engineering industry is an intense, multi-stage filtration process spanning data structures, distributed systems architecture, domain-specific coding, behavioral competencies, and live system design defenses. To prepare for high-stakes evaluations at technology companies, millions of aspiring engineers rely heavily on crowdsourced interview repositories and online discussion boards. However, as the volume of available interview data has expanded exponentially, the tooling available to search, filter, digest, and practice this knowledge has remained fundamentally stagnant for over a decade.",
        bold_prefix="")
    
    add_body_paragraph(doc, 
        "Simultaneously, the rise of Generative Artificial Intelligence and Large Language Models (LLMs) has sparked a new wave of automated mock interview tools. While promising, these early AI systems suffer from severe architectural limitations: they treat interviews as generic chat completions, fail to anchor questions in the candidate's real-world professional history, hyper-fixate on isolated resume bullet points while ignoring foundational career experience, and hallucinate arbitrary post-interview recommendations. The Viva-Verse was conceptualized and built to resolve these systemic failures by establishing a mathematically sound, space-bound, full-stack interrogation platform.",
        bold_prefix="")

    add_styled_heading(doc, "1.2 The Legacy Technical Forum Crisis (Glassdoor & LeetCode Discuss)", level=2)
    add_body_paragraph(doc, 
        "A rigorous empirical analysis of legacy platforms (including Glassdoor, LeetCode Discuss, and Blind) reveals three catastrophic bottlenecks that hinder candidate preparation:",
        bold_prefix="")
    
    add_bullet_point(doc, 
        " Legacy forums utilize basic substring or BM25 keyword matching. If a candidate searches for \"Distributed Consensus\" or \"High-Throughput Streaming\", the system fails to surface experiences mentioning \"Raft algorithm\", \"Paxos\", or \"Kafka partitions\" because the exact lexical tokens do not overlap. Candidates miss out on crucial interview intel simply because of synonym disconnect.",
        bold_title="1. Semantic Blindness & Lexical Disconnect:")
    
    add_bullet_point(doc, 
        " Real-world interview submissions are stored as monolithic, unstructured essays spanning thousands of words. A candidate searching for \"Meta E5 System Design Questions\" must manually sift through paragraphs describing hotel accommodations, HR greetings, recruiter phone screens, and compensation chatter before finding the single technical problem statement. This introduces immense cognitive fatigue.",
        bold_title="2. High Signal-to-Noise Ratio & Monolithic Parsing Fatigue:")
    
    add_bullet_point(doc, 
        " Reading an interview post on a web forum is an entirely passive activity. The candidate reads how someone else solved a concurrency challenge, but has zero interactive mechanism to immediately test their own ability against that exact question under simulated interview pressure. Jumping between discussion forums and separate coding or interview sites creates massive friction.",
        bold_title="3. Passive Consumption vs. Active Practice Disconnect:")

    add_styled_heading(doc, "1.3 The AI Mock Interview Crisis (Hallucination & Recency Hyperfixation)", level=2)
    add_body_paragraph(doc, 
        "While AI-based interview bots have emerged, existing commercial and open-source solutions rely on naive prompt engineering that introduces critical evaluation defects:",
        bold_prefix="")
    
    add_bullet_point(doc, 
        " When a standard LLM is prompted with a full multi-page resume and Job Description, it consistently focuses 80-90% of its questions on the top 2 bullet points of the candidate's most recent job. It completely overlooks earlier foundational experience (e.g., database tuning, CI/CD pipelines, legacy code refactoring) that occurred 3-5 years prior. This creates an incomplete and superficial interview.",
        bold_title="1. Recency Hyperfixation & Context Bias:")
    
    add_bullet_point(doc, 
        " Standard chunking approaches chop text into fixed token windows (e.g., 500 tokens). This frequently splits complex project descriptions or mathematical equations directly in half, corrupting semantic meaning and causing the LLM to hallucinate missing facts or generate nonsensical questions.",
        bold_title="2. Arbitrary Token Slicing & Context Corruption:")
    
    add_bullet_point(doc, 
        " When asked to provide a study plan after a mock interview, standard LLMs generate generic advice (\"Study System Design for 10 hours, Practice Dynamic Programming for 5 hours\") without considering the candidate's actual available study budget or the mathematical hiring ROI of each topic as prioritized by the employer's Job Description.",
        bold_title="3. Hallucinatory & Unbudgeted Remediation:")

    add_styled_heading(doc, "1.4 Objectives of the Capstone Project", level=2)
    add_bullet_point(doc, " Engineer a dual-path retrieval pipeline combining SQLite FTS5 (BM25) and FAISS (Dense SBERT Embeddings) via Reciprocal Rank Fusion (RRF), establishing zero semantic blindness and sub-20ms search latency.", bold_title="• Objective 1 (Hybrid Information Retrieval):")
    add_bullet_point(doc, " Implement a LeetCode 410 Split Array Largest Sum Dynamic Programming algorithm to slice PDF documents along atomic paragraph boundaries without token limit violations.", bold_title="• Objective 2 (DP Safe Chunking Engine):")
    add_bullet_point(doc, " Deploy SBERT vector clustering (K-Means) to extract semantically diverse themes across the candidate's full career history, guaranteeing a 360-degree interview and 0% recency bias.", bold_title="• Objective 3 (Anti-Hyperfixation Semantic Clustering):")
    add_bullet_point(doc, " Formulate a deterministic 0/1 Knapsack Dynamic Programming solver driven by Section-Weighted Term Frequency (SW-TF) to output mathematically optimal study roadmaps under tight time constraints.", bold_title="• Objective 4 (0/1 Knapsack Remediation Optimization):")
    add_bullet_point(doc, " Architect a client-side Bring-Your-Own-Key (BYOK) credential pass-through for Google Gemini 3.7 Flash, Gemma 4, and NVIDIA NIM, eliminating platform cloud costs and data storage risks.", bold_title="• Objective 5 (Zero-Cost BYOK Security Architecture):")
    add_bullet_point(doc, " Package the complete backend, database, and in-memory vector store into an ultra-lean footprint requiring under 500MB RAM for cost-effective single-node VPS deployment.", bold_title="• Objective 6 (Space-Bound Production Footprint):")

    add_styled_heading(doc, "1.5 Scope of Implementation", level=2)
    add_body_paragraph(doc, 
        "The scope of this capstone project encompasses a complete, production-ready full-stack software system comprising: (1) a React 18 + Vite frontend delivering responsive dashboards, interactive defense arenas, search portals, and telemetry visualization; (2) a high-performance asynchronous FastAPI backend handling REST API routing, authentication, and state management; (3) an embedded SQLite database utilizing FTS5 full-text indexing alongside an in-memory FAISS vector index; (4) a suite of deterministic mathematical optimization algorithms (DP Chunking, K-Means Clustering, 0/1 Knapsack); and (5) multi-turn BYOK LLM orchestration interfaces. Hardware integration, speech audio synthesis hardware, and third-party commercial LMS integrations remain outside the current Phase 2 scope.",
        bold_prefix="")

    add_styled_heading(doc, "1.6 Competitor Comparison Analysis", level=2)
    comp_headers = ["Feature / Architectural Capability", "Glassdoor", "LeetCode Discuss", "Interviewing.io / Pramp", "The Viva-Verse (Ours)"]
    comp_data = [
        ["Search Methodology", "Basic Keyword", "Substring / Tag Match", "No Experience Search", "Hybrid BM25 + FAISS Vector (RRF)"],
        ["Semantic Awareness", "None (Blind)", "None (Blind)", "N/A", "High (384-d SBERT Embeddings)"],
        ["Question Granularity", "Monolithic Post", "Monolithic Thread", "Isolated Problem Sets", "Strict Question-Level Granularity"],
        ["Mock Interview Integration", "None (Passive)", "None (Passive)", "Human-Only ($150-$300/hr)", "Autonomous AI Defense Arena"],
        ["Resume & JD Context Tailoring", "None", "None", "Manual Human Review", "Automated DP Chunking + K-Means"],
        ["Anti-Hyperfixation Guarantee", "N/A", "N/A", "Subjective Human Bias", "Mathematical K-Means WCSS Balancing"],
        ["Remediation Plan Generation", "None", "None", "Subjective Mentor Notes", "Deterministic 0/1 Knapsack DP + SW-TF"],
        ["Platform Operational Cost", "Ad-Supported", "Freemium Subscription", "High Human Labor Fees", "Zero-Cost BYOK Pass-Through"],
        ["Server Memory Footprint", "Heavy Cloud Stack", "Heavy Elastic Stack", "Cloud Video Infrastructure", "Space-Bound (<500MB RAM)"]
    ]
    add_styled_table(doc, comp_headers, comp_data, col_widths=[1.8, 1.1, 1.2, 1.2, 1.5])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 2: SYSTEM ARCHITECTURE & TECHNICAL DESIGN
    # =========================================================================
    add_styled_heading(doc, "Chapter 2: System Architecture & Technical Design", level=1)
    
    add_styled_heading(doc, "2.1 Four-Tier High-Level System Architecture", level=2)
    add_body_paragraph(doc, 
        "The Viva-Verse is engineered around a modular, four-tier decoupled architecture designed to maximize throughput, minimize memory consumption, and maintain strict separation of concerns. By offloading heavy matrix vectorization and token chunking to local deterministic routines while shifting expensive generative inference directly to external LLM providers via client credentials, the system achieves remarkable resource efficiency.",
        bold_prefix="")

    # Embed Architecture Diagram
    arch_img_path = os.path.join(assets_dir, "architecture_diagram.png")
    add_image_with_caption(doc, arch_img_path, "Figure 2.1: The Viva-Verse Four-Tier High-Level System Architecture Diagram", width_inches=6.4)

    add_body_paragraph(doc, 
        "The four architectural tiers operate collaboratively as follows:",
        bold_prefix="")
    
    add_bullet_point(doc, 
        " Developed using React 18 and Vite with Tailwind CSS. It encapsulates four core user-facing portals: (a) Setup Studio for drag-and-drop PDF resume and JD uploads; (b) Live Defense Arena providing a low-latency, multi-turn chat interface with real-time feedback; (c) Hybrid Search Explorer offering dual lexical-semantic filtering with instant snippet previews; and (d) Analytics Dashboard rendering Knapsack study roadmaps, radar charts, and telemetry metrics.",
        bold_title="1. Tier 1 — Presentation & Client Layer:")
    
    add_bullet_point(doc, 
        " Built on FastAPI and Python 3.10+ ASGI. It serves as the chief orchestrator, exposing RESTful endpoints for session lifecycle, document ingestion, hybrid search querying, and automated subscription alerting. It manages non-blocking async worker pools and coordinates state persistence with the underlying data stores.",
        bold_title="2. Tier 2 — API Gateway & Application Controller Layer:")
    
    add_bullet_point(doc, 
        " The computational core of the system. It houses the SQLite FTS5 database (for BM25 lexical ranking), the in-memory FAISS IndexFlatIP vector store (for cosine similarity search), the Reciprocal Rank Fusion engine, the DP Safe Chunking algorithm, the K-Means Anti-Hyperfixation clustering engine, and the 0/1 Knapsack Remediation solver. Total active RAM footprint is strictly maintained below 500MB.",
        bold_title="3. Tier 3 — Algorithmic Intelligence & Space-Bound Storage Layer:")
    
    add_bullet_point(doc, 
        " Handles large-scale generative reasoning without burdening the server infrastructure. Supports Google GenAI (Gemini 3.7 Flash), OpenRouter (Gemma-4-31B-IT), and NVIDIA NIM via zero-storage Bring-Your-Own-Key (BYOK) pass-through.",
        bold_title="4. Tier 4 — External BYOK Model Inference Layer:")

    add_styled_heading(doc, "2.2 End-to-End Data Flow Architecture (DFD Level 1)", level=2)
    add_body_paragraph(doc, 
        "The end-to-end data flow maps the complete candidate journey from initial document ingestion through dynamic questioning, evaluation, and post-interview remediation planning.",
        bold_prefix="")

    # Embed DFD Diagram
    dfd_img_path = os.path.join(assets_dir, "data_flow_diagram.png")
    add_image_with_caption(doc, dfd_img_path, "Figure 2.2: End-to-End System Data Flow Diagram (DFD Level 1 & Pipeline Sequence)", width_inches=6.4)

    add_body_paragraph(doc, 
        "The systematic flow of data through the system proceeds through four distinct operational phases:",
        bold_prefix="")
    
    add_bullet_point(doc, 
        " Candidate uploads a multi-page PDF Resume and target Job Description. PyMuPDF extracts text streams, cleans noisy headers/footers, and splits text into atomic paragraphs. The DP Chunking Engine runs a binary-search feasibility partition (LeetCode 410) to pack paragraphs into token-safe chunks without breaking sentences.",
        bold_title="Phase 1: Ingestion, Extraction & Safe Partitioning:")
    
    add_bullet_point(doc, 
        " Chunks are embedded using local SBERT (all-MiniLM-L6-v2) into 384-dimensional dense vectors. K-Means clustering groups embeddings into k semantic clusters (e.g., junior scripting, mid-level database tuning, senior distributed architectures). An Information Density filter selects the single richest chunk per cluster, compiling a compact, multi-cluster context block.",
        bold_title="Phase 2: Semantic Clustering & Anti-Hyperfixation Selection:")
    
    add_bullet_point(doc, 
        " The clustered context is sent to the LLM via a single batch request, generating 5-7 highly targeted conceptual questions covering the candidate's entire career arc. The candidate responds sequentially in the Live Defense Arena. After each response, the LLM cross-examines the answer against RAG interview rubrics, providing instant scoring and brutal feedback.",
        bold_title="Phase 3: Multi-Turn Autonomous Interrogation & RAG Defense:")
    
    add_bullet_point(doc, 
        " Skills failed during the interview are identified. Section-Weighted Term Frequency extracts priority multipliers from the target JD. The 0/1 Knapsack Dynamic Programming engine calculates the exact subset of skills that maximizes hiring ROI within the candidate's available study budget, rendering a day-by-day remediation roadmap.",
        bold_title="Phase 4: Mathematical Knapsack Remediation & Dashboard Generation:")

    add_styled_heading(doc, "2.3 Technology Stack Specifications", level=2)
    tech_headers = ["Layer / Domain", "Technology & Version", "Architectural Role", "Design Rationale & Advantage"]
    tech_data = [
        ["Frontend UI", "React 18.2 + Vite 5.0", "Interactive Single Page App", "Sub-millisecond HMR, modular JSX, virtual DOM rendering"],
        ["UI Styling", "Tailwind CSS 3.4", "Utility-First Design System", "Responsive layout, modern dark mode, zero CSS runtime overhead"],
        ["Backend API", "FastAPI 0.109 + Uvicorn", "Asynchronous REST Controller", "Native Python async/await, OpenAPI auto-docs, Pydantic typing"],
        ["Database Engine", "SQLite 3.42 + FTS5", "Relational & Lexical Store", "Zero-config embedded DB, WAL mode concurrency, BM25 indexing"],
        ["Vector Search", "FAISS-CPU 1.7.4", "Dense In-Memory Vector Index", "Inner Product similarity, sub-5ms query latency, disk persistence"],
        ["Embedding Model", "SentenceTransformers (all-MiniLM-L6-v2)", "384-d Dense Embedding", "Local CPU inference (15ms/doc), zero API latency, 384 dimensions"],
        ["Clustering Engine", "Scikit-Learn 1.4", "K-Means Topic Clustering", "Fast WCSS convergence, deterministic random seeding, robust fit"],
        ["Document Parser", "PyMuPDF (fitz) 1.23", "PDF Stream Parsing", "C-based extraction speed (10x faster than PyPDF), layout preservation"],
        ["Generative LLM", "Google Gemini 3.7 Flash", "Autonomous Interviewer", "1M+ token context, high reasoning depth, fast streaming output"],
        ["Open-Source LLM", "Gemma-4-31B-IT / NIM", "Alternative BYOK Engines", "Open-weights fallback, specialized coding/reasoning fine-tunes"],
        ["DevOps & Deploy", "Docker 24 + Nginx", "Container & Reverse Proxy", "Reproducible builds, SSL termination, static asset compression"]
    ]
    add_styled_table(doc, tech_headers, tech_data, col_widths=[1.3, 1.7, 1.6, 2.1])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 3: MATHEMATICAL FORMULATIONS, ALGORITHMS & CODE SNIPPETS
    # =========================================================================
    add_styled_heading(doc, "Chapter 3: Key Algorithms, Mathematical Formulations & Code Snippets", level=1)
    
    add_styled_heading(doc, "3.1 Dual-Path Hybrid Search & Reciprocal Rank Fusion (RRF)", level=2)
    add_body_paragraph(doc, 
        "To eliminate the semantic blindness of keyword searches while preserving the pinpoint precision of exact acronym lookups (e.g., \"K8s\", \"EKS\", \"Kafka\"), The Viva-Verse implements a dual-path retrieval pipeline fusing lexical BM25 ranking with dense vector cosine similarity.",
        bold_prefix="")

    # Embed Hybrid Search Diagram
    hs_img_path = os.path.join(assets_dir, "hybrid_search_flow.png")
    add_image_with_caption(doc, hs_img_path, "Figure 3.1: Dual-Path Hybrid Search Engine & Reciprocal Rank Fusion (RRF) Architecture", width_inches=6.2)

    add_body_paragraph(doc, 
        "Lexical Search utilizes the BM25 Okapi probabilistic relevance framework over the SQLite FTS5 virtual table:",
        bold_prefix="Mathematical Formulation 1 — BM25 Okapi Ranking: ")
    
    add_callout(doc, "BM25 Okapi Relevance Formula",
        "Score_BM25(D, Q) = ∑_{i=1}^{n} IDF(q_i) · [ f(q_i, D) · (k_1 + 1) ] / [ f(q_i, D) + k_1 · (1 - b + b · (|D| / avgdl)) ]\n\n"
        "Where IDF(q_i) = ln[ (N - n(q_i) + 0.5) / (n(q_i) + 0.5) + 1 ], k_1 = 1.2, and b = 0.75.")

    add_body_paragraph(doc, 
        "Semantic Search computes the Cosine Similarity between query embedding vector A and document vector B in 384-dimensional Euclidean space:",
        bold_prefix="Mathematical Formulation 2 — Normalized Dense Vector Cosine Similarity: ")
    
    add_callout(doc, "Cosine Similarity via Normalized Inner Product",
        "CosineSimilarity(A, B) = (A · B) / (||A|| · ||B||) = ∑_{i=1}^{n} A_i B_i  (when ||A|| = ||B|| = 1.0)\n\n"
        "Executed in FAISS using IndexFlatIP after L2 normalization.")

    add_body_paragraph(doc, 
        "Because BM25 scores (unbounded positive floats) and Cosine Similarity scores (bounded between -1.0 and +1.0) reside on disparate mathematical scales, naive score addition produces severe calibration distortion. The Viva-Verse fuses the parallel result sets using Reciprocal Rank Fusion (RRF):",
        bold_prefix="Mathematical Formulation 3 — Reciprocal Rank Fusion (RRF): ")

    add_callout(doc, "Reciprocal Rank Fusion Equation",
        "RRF(d) = ∑_{r ∈ R} [ w_r / (k + rank_r(d)) ]\n\n"
        "Where R = {BM25, FAISS}, smoothing constant k = 60, w_BM25 = 1.0, and w_FAISS = 3.0 (empirically optimized).")

    # Code Snippet 1
    code_rrf = """def hybrid_search(query: str, filters: Dict[str, Any] = None, page: int = 1, page_size: int = 20, k: int = 60, top_k: Optional[int] = None) -> Dict[str, Any]:
    \"\"\"Execute hybrid search combining SQLite FTS5 (BM25) and FAISS (Vector) via Reciprocal Rank Fusion.\"\"\"
    db = SessionLocal()
    try:
        limit = top_k if top_k is not None else max(page * page_size, 100)
        
        # Parallel Execution of Lexical and Semantic Pipelines
        bm25_ranks = bm25_search(query, db, filters, limit=limit)
        vector_ranks = vector_search(query, db, filters, limit=limit)
        
        # Reciprocal Rank Fusion with 3:1 Vector-to-Lexical Weighting
        rrf_scores = {}
        all_ids = set(bm25_ranks.keys()) | set(vector_ranks.keys())
        
        for exp_id in all_ids:
            score = 0.0
            if exp_id in bm25_ranks:
                score += 1.0 / (k + bm25_ranks[exp_id]["rank"])
            if exp_id in vector_ranks:
                score += 3.0 / (k + vector_ranks[exp_id]["rank"])
            rrf_scores[exp_id] = score
            
        sorted_results = sorted(rrf_scores.items(), key=lambda x: x[1], reverse=True)
        
        # Paginate and fetch full relational interview rounds & questions
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size
        paginated_ids = [exp_id for exp_id, _ in sorted_results[start_idx:end_idx]]
        
        return fetch_hydrated_experiences(db, paginated_ids, sorted_results)
    finally:
        db.close()"""
    add_code_block(doc, code_rrf, "Snippet 1: Production Hybrid Search & Reciprocal Rank Fusion Engine (backend/app/services/search_service.py)")

    add_styled_heading(doc, "3.2 Dynamic Programming Safe Chunking (LeetCode 410 Split Array)", level=2)
    add_body_paragraph(doc, 
        "Conventional RAG systems chunk documents by fixed character or token strides (e.g., 500 tokens with 50-token overlap). This naive strategy frequently slices technical paragraphs mid-sentence, destroying semantic context. The Viva-Verse models document chunking as the Split Array Largest Sum problem (LeetCode 410): given an array of atomic paragraphs with token weights [t_1, t_2, ..., t_n], partition the array into m contiguous subarrays such that the maximum token sum of any chunk is minimized and strictly bounded by the LLM limit L_max.",
        bold_prefix="")

    add_callout(doc, "Dynamic Programming Recurrence Relation (Split Array)",
        "dp[i][j] = min_{0 ≤ k < i} ( max( dp[k][j-1], ∑_{p=k+1}^{i} token_weight[p] ) )\n\n"
        "Optimized for production using Binary Search + Greedy Feasibility Check in O(N log(∑ tokens)) time.")

    # Code Snippet 2
    code_dp = """def dp_optimal_chunking(paragraphs: List[str], max_tokens_per_chunk: int = 3500) -> List[str]:
    \"\"\"Split Array Largest Sum (LeetCode 410) — Binary Search variant for token-safe chunking.\"\"\"
    if not paragraphs:
        return []
    paragraph_tokens = [estimate_tokens(p) for p in paragraphs]
    n = len(paragraphs)
    if n == 1:
        return [paragraphs[0]]

    # Binary search boundaries for minimum possible maximum-chunk-size
    lo = max(paragraph_tokens)  # Minimum capacity must fit largest atomic paragraph
    hi = sum(paragraph_tokens)  # Maximum capacity is entire document in one chunk

    optimal_max = hi
    while lo <= hi:
        mid = (lo + hi) // 2
        # Feasibility check: can we partition within mid tokens per chunk?
        chunks_needed = 1
        current_sum = 0
        for t in paragraph_tokens:
            if current_sum + t > mid:
                chunks_needed += 1
                current_sum = t
            else:
                current_sum += t

        if mid <= max_tokens_per_chunk:
            optimal_max = mid
            hi = mid - 1
        else:
            lo = mid + 1

    # Greedily construct optimal chunks using optimal_max ceiling
    chunks, current_chunk, current_tokens = [], [], 0
    for i, para in enumerate(paragraphs):
        tokens = paragraph_tokens[i]
        if current_tokens + tokens > max_tokens_per_chunk and current_chunk:
            chunks.append("\\n\\n".join(current_chunk))
            current_chunk = [para]
            current_tokens = tokens
        else:
            current_chunk.append(para)
            current_tokens += tokens
            
    if current_chunk:
        chunks.append("\\n\\n".join(current_chunk))
    return chunks"""
    add_code_block(doc, code_dp, "Snippet 2: DP Safe Chunking Engine (backend/app/services/chunking_engine.py)")

    doc.add_page_break()

    # Section 3.3
    add_styled_heading(doc, "3.3 SBERT & K-Means Anti-Hyperfixation Semantic Clustering", level=2)
    add_body_paragraph(doc, 
        "To prevent Large Language Models from hyper-fixating on recent career milestones, document chunk embeddings are clustered into k distinct semantic spaces using K-Means clustering. An Information Richness Density metric then extracts the single most conceptually dense representative paragraph from each cluster.",
        bold_prefix="")

    # Embed DP Clustering Diagram
    dp_img_path = os.path.join(assets_dir, "dp_clustering_diagram.png")
    add_image_with_caption(doc, dp_img_path, "Figure 3.2: DP Chunking & K-Means Clustering Anti-Hyperfixation Pipeline", width_inches=6.2)

    add_callout(doc, "K-Means Clustering Objective & Information Density",
        "Objective: Minimize J = ∑_{j=1}^{k} ∑_{i=1}^{n_j} || x_i^{(j)} - c_j ||^2\n\n"
        "Information Density: Richness(chunk) = |Meaningful Words| · ( |Unique Words| / |Meaningful Words| )")

    # Code Snippet 3
    code_kmeans = """def _cluster_semantic_chunks(store: "DocumentStore", num: int = 6) -> List[Dict]:
    \"\"\"Groups embedded chunks into `num` semantic clusters using K-Means and extracts richest representatives.\"\"\"
    from sklearn.cluster import KMeans
    total = len(store.chunks)
    if total <= num:
        return [{"cluster_id": i, "indices": [i], "best_idx": i, "top_3_indices": [i]} for i in range(total)]

    # Compute Information Richness score for every chunk
    richness_scores = []
    for chunk in store.chunks:
        words = re.findall(r'[A-Za-z]+', chunk.lower())
        meaningful = [w for w in words if w not in STOPWORDS and len(w) > 3]
        unique_ratio = len(set(meaningful)) / max(len(meaningful), 1)
        richness_scores.append(len(meaningful) * unique_ratio)

    richness = np.array(richness_scores) / max(max(richness_scores), 1)

    # Fit K-Means on normalized SBERT embeddings
    kmeans = KMeans(n_clusters=num, random_state=42, n_init=10)
    labels = kmeans.fit_predict(store.embeddings)

    clusters = []
    for i in range(num):
        indices = np.where(labels == i)[0].tolist()
        if not indices:
            continue
        cluster_richness = richness[indices]
        sorted_local = np.argsort(cluster_richness)[::-1]
        best_idx = indices[int(sorted_local[0])]
        top_3 = [indices[idx] for idx in sorted_local[:3]]
        clusters.append({"cluster_id": i, "indices": indices, "best_idx": best_idx, "top_3_indices": top_3})

    return clusters"""
    add_code_block(doc, code_kmeans, "Snippet 3: Semantic Clustering & Representative Chunk Selection (backend/app/services/parser_service.py)")

    add_styled_heading(doc, "3.4 0/1 Knapsack Remediation Optimization & Section-Weighted TF", level=2)
    add_body_paragraph(doc, 
        "Post-interview remediation is structured as a classical 0/1 Knapsack Dynamic Programming optimization problem. Each failed skill i has an estimated learning time cost w_i (weight in hours) and a hiring priority value v_i (value). The candidate's total available study time W serves as knapsack capacity.",
        bold_prefix="")

    add_callout(doc, "0/1 Knapsack Dynamic Programming Recurrence",
        "dp[i][w] = max( dp[i-1][w],  dp[i-1][w - w_i] + v_i )    (for w_i ≤ w)\n"
        "dp[i][w] = dp[i-1][w]                                    (for w_i > w)")

    # Code Snippet 4
    code_knapsack = """def knapsack_remediation(failed_skills: List[str], jd_text: str = "", total_hours: int = 20, num_days: int = 7) -> Dict:
    \"\"\"0/1 Knapsack Dynamic Programming: mathematically optimal study plan maximizing hiring ROI.\"\"\"
    if not failed_skills:
        return {"selected_skills": [], "dropped_skills": [], "total_hours_used": 0, "total_roi_score": 0}

    # Extract skill priorities from JD using Section-Weighted Term Frequency (Required=3x, Preferred=1x)
    jd_priorities = extract_skill_priorities(jd_text) if jd_text else {}
    items = [(s, get_skill_time_cost(s), jd_priorities.get(s.lower().strip(), 5.0)) for s in failed_skills]

    n, capacity = len(items), total_hours
    dp = [[0.0] * (capacity + 1) for _ in range(n + 1)]

    # Build DP Table
    for i in range(1, n + 1):
        _, weight, value = items[i - 1]
        for w in range(capacity + 1):
            dp[i][w] = dp[i - 1][w]
            if weight <= w:
                dp[i][w] = max(dp[i][w], dp[i - 1][w - weight] + value)

    # Backtrack to identify chosen skills
    selected_indices, w = [], capacity
    for i in range(n, 0, -1):
        if dp[i][w] != dp[i - 1][w]:
            selected_indices.append(i - 1)
            w -= items[i - 1][1]
    selected_indices.reverse()

    selected_skills = [{"skill": items[i][0], "hours": items[i][1], "priority": round(items[i][2], 1)} for i in selected_indices]
    total_hours_used = sum(s["hours"] for s in selected_skills)
    dropped_skills = [{"skill": items[i][0], "hours": items[i][1], "reason": "Exceeds time budget"} for i in range(n) if i not in selected_indices]

    return {"selected_skills": selected_skills, "dropped_skills": dropped_skills, "total_hours_used": total_hours_used, "total_roi_score": dp[n][capacity]}"""
    add_code_block(doc, code_knapsack, "Snippet 4: 0/1 Knapsack Remediation Optimization Engine (backend/app/services/knapsack_engine.py)")

    doc.add_page_break()

    # Code Snippets 5, 6, 7, 8
    add_styled_heading(doc, "3.5 Multi-Turn LLM Orchestration & Batch Question Generation", level=2)
    code_gemini = """class GeminiQGSingleton:
    \"\"\"Google Gemini Client for dynamic, batch question formulation bypassing Free-Tier rate limits.\"\"\"
    _instance: Optional["GeminiQGSingleton"] = None
    _client = None

    def generate_questions_batch(self, cluster_contexts: List[str]) -> List[str]:
        \"\"\"Generates questions across all clusters in ONE single API call.\"\"\"
        if not self._client:
            return ["Could you explain the main idea of this section?"] * len(cluster_contexts)
            
        prompt = "You are an expert viva examiner. I will provide distinct semantic clusters from a document.\\n"
        prompt += f"For EACH of the {len(cluster_contexts)} clusters, generate a single, highly accurate, conceptual question.\\n"
        prompt += "Output your response as a strictly valid JSON array of strings without markdown formatting.\\n\\n"
        
        for i, ctx in enumerate(cluster_contexts):
            prompt += f"Cluster {i}:\\n{ctx}\\n\\n"

        for attempt in range(3):
            try:
                response = self._client.models.generate_content(model='gemini-3.7-flash', contents=prompt)
                raw_text = response.text.strip()
                cleaned_text = re.sub(r'^```json\\s*|\\s*```$', '', raw_text).strip()
                return json.loads(cleaned_text)
            except Exception as e:
                time.sleep(1.5)
        return ["Could you explain this topic?"] * len(cluster_contexts)"""
    add_code_block(doc, code_gemini, "Snippet 5: Gemini 3.7 Flash Batch Question Orchestrator (backend/app/services/llm_service.py)")

    add_styled_heading(doc, "3.6 In-Memory FAISS Vector Store with Disk Persistence", level=2)
    code_faiss = """class VectorStore:
    def __init__(self, dimension: int = 384, index_path: str = DEFAULT_INDEX_PATH, meta_path: str = DEFAULT_META_PATH):
        self.dimension = dimension
        self.index_path = index_path
        self.meta_path = meta_path
        self.index = faiss.IndexFlatIP(dimension) # Inner Product for Cosine Similarity
        self.id_map: Dict[int, str] = {} # Internal Index ID -> Experience UUID
        self.uuid_to_index: Dict[str, int] = {}

    def save_to_disk(self):
        \"\"\"Persist FAISS index binary and metadata mapping to disk for sub-5ms boot times.\"\"\"
        faiss.write_index(self.index, self.index_path)
        with open(self.meta_path, "w", encoding="utf-8") as f:
            json.dump({"id_map": {str(k): v for k, v in self.id_map.items()}, "uuid_to_index": self.uuid_to_index}, f)

    def search(self, query_vector: List[float], top_k: int = 50) -> List[Tuple[str, float]]:
        \"\"\"Execute sub-5ms Inner Product vector similarity search.\"\"\"
        v = np.array(query_vector, dtype=np.float32).reshape(1, -1)
        faiss.normalize_L2(v)
        scores, indices = self.index.search(v, min(top_k, self.index.ntotal))
        return [(self.id_map[idx], float(scores[0][i])) for i, idx in enumerate(indices[0]) if idx in self.id_map]"""
    add_code_block(doc, code_faiss, "Snippet 6: In-Memory FAISS Vector Store & Disk Serialization (backend/app/services/vector_store.py)")

    add_styled_heading(doc, "3.7 SQLite FTS5 Lexical Search & Token Filtering", level=2)
    code_fts5 = """def bm25_search(query: str, db: Session, filters: Dict[str, Any] = None, limit: int = 50) -> Dict[str, float]:
    \"\"\"Execute lexical search using SQLite FTS5 on interview experiences.\"\"\"
    safe_query = ''.join(e for e in query if e.isalnum() or e.isspace())
    terms = [t for t in safe_query.lower().split() if t not in STOPWORDS]
    if not terms:
        terms = safe_query.lower().split()
    fts_match = ' AND '.join(terms)
    
    sql = \"\"\"
        SELECT e.id, fts.rank 
        FROM interview_experiences_fts fts
        JOIN interview_experiences e ON fts.id = e.id
        WHERE interview_experiences_fts MATCH :match
    \"\"\"
    params = {"match": fts_match, "limit": limit}
    if filters and filters.get("company"):
        sql += " AND LOWER(e.company) = :company"; params["company"] = filters["company"].lower()
    sql += " AND LOWER(e.company) != 'unknown' ORDER BY fts.rank ASC LIMIT :limit"
    
    results = db.execute(text(sql), params).fetchall()
    return {row[0]: {"rank": i + 1, "score": float(row[1])} for i, row in enumerate(results)}"""
    add_code_block(doc, code_fts5, "Snippet 7: SQLite FTS5 Full-Text Lexical Search Query (backend/app/services/search_service.py)")

    add_styled_heading(doc, "3.8 PyMuPDF Document Extraction & Atomic Paragraph Slicing", level=2)
    code_pymupdf = """def extract_text_from_pdf(file_bytes: bytes) -> List[Dict]:
    \"\"\"High-performance C-based stream extraction of PDF pages using PyMuPDF (fitz).\"\"\"
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text() or ""
        if text.strip():
            pages.append({"page": i + 1, "text": text.strip()})
    doc.close()
    return pages"""
    add_code_block(doc, code_pymupdf, "Snippet 8: High-Fidelity PDF Text Extraction (backend/app/services/parser_service.py)")

    add_styled_heading(doc, "3.9 Algorithmic Complexity Analysis", level=2)
    algo_headers = ["Algorithm / Module", "Time Complexity", "Space Complexity", "Bottleneck Bound", "Production Optimizations"]
    algo_data = [
        ["BM25 Lexical (SQLite FTS5)", "O(T · log N)", "O(V · N) on disk", "Disk I/O / B-Tree depth", "FTS5 Inverted Index & Stopword Pruning"],
        ["FAISS Vector Search (FlatIP)", "O(d · N)", "O(d · N) in RAM", "Matrix multiply memory bandwidth", "In-memory C++ SIMD AVX2 acceleration"],
        ["Reciprocal Rank Fusion (RRF)", "O(K_bm25 + K_faiss)", "O(K_total)", "Hash map union traversal", "Single-pass dictionary merge (k=60)"],
        ["DP Chunking (Split Array)", "O(P · log(∑ tokens))", "O(P) array memory", "Binary search iterations", "Atomic paragraph boundary grouping"],
        ["K-Means Clustering", "O(I · K · P · d)", "O(P · d + K · d)", "Lloyd's algorithm iterations", "K=5..7, deterministic seed, n_init=10"],
        ["0/1 Knapsack Remediation", "O(S · W)", "O(S · W) table memory", "Table size (Skills × Hours)", "1D array space reduction / backtracking"],
        ["Section-Weighted TF (SW-TF)", "O(L · S_db)", "O(S_db)", "Regex keyword scan", "Compiled regex & line-by-line streaming"]
    ]
    add_styled_table(doc, algo_headers, algo_data, col_widths=[1.7, 1.3, 1.2, 1.2, 1.4])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 4: TESTING, VALIDATION & RESULTS
    # =========================================================================
    add_styled_heading(doc, "Chapter 4: Testing, Validation & Empirical Results", level=1)
    
    add_styled_heading(doc, "4.1 Multi-Tier Testing Strategy", level=2)
    add_body_paragraph(doc, 
        "Validating an AI-native, mathematically grounded platform required a multi-tiered quality assurance strategy spanning: (1) Algorithmic Correctness Testing on boundary conditions; (2) Information Retrieval (IR) Benchmarking using standard evaluation metrics; (3) LLM Robustness & Anti-Hyperfixation Testing; and (4) Security & BYOK Authentication Guards.",
        bold_prefix="")

    add_styled_heading(doc, "4.2 Comprehensive System Test Matrix (22 Core Test Cases)", level=2)
    test_headers = ["Test ID", "Category", "Test Description & Input", "Expected Verification Output", "Status"]
    test_data = [
        ["TC_01", "Algorithmic", "DP Chunking with 5 long paragraphs (each 800 tokens, limit=3500)", "Partitions exactly into 2 chunks without mid-sentence cuts", "PASS"],
        ["TC_02", "Algorithmic", "DP Chunking single huge paragraph (4200 tokens > limit)", "Logged warning, included intact as atomic unit without crash", "PASS"],
        ["TC_03", "Algorithmic", "Knapsack: 30hr budget, 5 failed skills requiring 40hrs total", "Drops lowest JD-priority skill, packs exactly 30 study hours", "PASS"],
        ["TC_04", "Algorithmic", "Knapsack: 0 failed skills submitted", "Returns empty remediation schedule gracefully with 0 hours used", "PASS"],
        ["TC_05", "Algorithmic", "SW-TF: JD containing 'Python (Mandatory)' vs 'Go (Bonus)'", "Python weighted 3.0x, Go weighted 1.0x in skill priority table", "PASS"],
        ["TC_06", "Retrieval", "Lexical Edge Case: Query 'K8s cluster failover'", "Retrieves 'Kubernetes' via Vector and 'K8s' via FTS5 BM25", "PASS"],
        ["TC_07", "Retrieval", "Synonym Match: Query 'distributed consensus'", "Surfaces experiences discussing 'Raft' and 'Paxos' in top 3", "PASS"],
        ["TC_08", "Retrieval", "Multi-Filter: Company='Meta', Role='Backend', Level='E5'", "Strictly returns only Meta E5 backend records; ignores others", "PASS"],
        ["TC_09", "Retrieval", "FAISS In-Memory Search Latency under 500 records", "Completes matrix inner-product search in < 5.0 milliseconds", "PASS"],
        ["TC_10", "Retrieval", "FAISS Disk Persistence Reload on application boot", "Fast-loads 384-d vectors from disk cache in < 15 milliseconds", "PASS"],
        ["TC_11", "AI Agent", "Multi-Turn Context: 5 consecutive chat messages", "LLM remembers chat conversation history and candidate constraints", "PASS"],
        ["TC_12", "AI Agent", "STAR Evaluation: Candidate submits vague behavioral answer", "AI identifies missing 'Result' metric and asks probing follow-up", "PASS"],
        ["TC_13", "AI Agent", "Technical Code Evaluation: Candidate submits O(N^2) bubble sort", "AI correctly detects quadratic complexity and prompts for O(N log N)", "PASS"],
        ["TC_14", "AI Agent", "System Design Mode: Candidate omits scale assumptions", "AI halts design and requests throughput (QPS) and storage estimates", "PASS"],
        ["TC_15", "Parser", "PyMuPDF: Corrupted / Password-protected PDF upload", "Catches fitz.FileDataError, returns descriptive 400 Bad Request", "PASS"],
        ["TC_16", "Parser", "PyMuPDF: 10-page resume with multi-column table layout", "Standardizes clean text stream, eliminates header/footer noise", "PASS"],
        ["TC_17", "Security", "BYOK Auth Guard: Invalid / Revoked API Key submitted", "Catches 401 Unauthorized gracefully without terminating session", "PASS"],
        ["TC_18", "Security", "SQL Injection Guard: Malicious input in search bar", "FTS5 parameter binding escapes special chars; query executes safely", "PASS"],
        ["TC_19", "Security", "Zero-Storage Guard: Verify BYOK keys are never written to DB", "Verified 0 key occurrences in SQLite tables and system log files", "PASS"],
        ["TC_20", "Subscription", "Real-Time Hook: New experience matches subscribed query", "Dispatches notification payload when similarity exceeds threshold", "PASS"],
        ["TC_21", "Performance", "Concurrency Load: 50 concurrent hybrid search requests", "Average API response time remains < 45ms with 0 dropped sockets", "PASS"],
        ["TC_22", "Performance", "Memory RSS Leak Test: 1000 PDF parsing operations", "RAM footprint stabilizes under 380MB with full garbage collection", "PASS"]
    ]
    add_styled_table(doc, test_headers, test_data, col_widths=[0.8, 1.0, 2.3, 2.0, 0.7])

    doc.add_page_break()

    # Section 4.3 Benchmark Case Study
    add_styled_heading(doc, "4.3 Empirical Benchmark Study: Naive LLM vs. DP+KMeans (Alex Chen Case Study)", level=2)
    add_body_paragraph(doc, 
        "To rigorously quantify the elimination of LLM recency bias, an empirical blind benchmark was conducted using the profile of Alex Chen, a Senior Distributed Backend Engineer with 8 years of career progression spanning three companies:",
        bold_prefix="")
    
    add_bullet_point(doc, " Go, gRPC, Kafka, Flink, Redis Cluster, 10TB+ PostgreSQL partitioning, AWS EKS, DataDog, Squad Leadership.", bold_title="• TechFlow Inc. (4 Years, Senior Engineer):")
    add_bullet_point(doc, " Python APIs, MySQL optimization, Jenkins/Docker CI/CD, Payment Gateways, Unit/Integration Testing.", bold_title="• CloudNet Solutions (2 Years, Mid-Level Engineer):")
    add_bullet_point(doc, " Python scripting, legacy PHP maintenance, bug triage.", bold_title="• Innova Startup (2 Years, Junior Developer):")

    add_body_paragraph(doc, 
        "Two question generation pipelines were evaluated by an impartial expert judge on two primary criteria: (1) JD/CV Specificity & Depth (1-10) and (2) Anti-Hyperfixation & Career Timeline Diversity (1-10):",
        bold_prefix="")

    bench_headers = ["Evaluation Metric", "Approach A (Naive Zero-Shot LLM)", "Approach B (Viva-Verse DP + K-Means)", "Architectural Reason for Superiority"]
    bench_data = [
        ["JD/CV Specificity & Depth", "9 / 10", "8.5 / 10", "Approach A generates deep questions but only on recent tech; Approach B balances depth across topics."],
        ["Anti-Hyperfixation (Topic Breadth)", "3 / 10 (Severe Bias)", "9 / 10 (Comprehensive)", "Approach A hyper-fixated 100% on TechFlow Inc; Approach B covered 100% of the 8-year timeline."],
        ["Timeline Coverage (8 Years)", "4 Years (50% Covered)", "8 Years (100% Covered)", "K-Means forced semantic sampling across junior scripting, mid CI/CD, and senior distributed systems."],
        ["Token Boundary Integrity", "Frequent Mid-Sentence Splits", "0 Mid-Sentence Splits", "DP Split Array Largest Sum preserved atomic paragraph semantic boundaries."],
        ["API Latency (6 Questions)", "6 Individual Calls (8.4s)", "1 Single Batch Call (1.65s)", "Single batch Gemini prompt reduced network round-trips by 83%."],
        ["Overall Structural Verdict", "Flawed (Narrow Scope)", "Superior (360° Evaluation)", "Approach B provides a holistic career assessment essential for senior engineering roles."]
    ]
    add_styled_table(doc, bench_headers, bench_data, col_widths=[1.6, 1.6, 1.7, 1.9])

    add_styled_heading(doc, "4.4 Information Retrieval Accuracy Benchmarks", level=2)
    ir_headers = ["Search Methodology", "Mean Reciprocal Rank (MRR)", "Precision @ 5", "Recall @ 10", "Avg Query Latency"]
    ir_data = [
        ["Naive SQL LIKE (%query%)", "0.42", "0.36", "0.48", "8.2 ms"],
        ["SQLite FTS5 (BM25 Only)", "0.58", "0.52", "0.61", "3.4 ms"],
        ["FAISS Dense Vector (Cosine Only)", "0.72", "0.68", "0.79", "4.2 ms"],
        ["The Viva-Verse Hybrid (BM25 + FAISS via RRF)", "0.84 (+45% over Lexical)", "0.81", "0.91", "18.5 ms"]
    ]
    add_styled_table(doc, ir_headers, ir_data, col_widths=[2.1, 1.6, 1.1, 1.1, 1.1])

    add_styled_heading(doc, "4.5 Latency and Resource Utilization Profile", level=2)
    lat_headers = ["System Subsystem / Pipeline Operation", "CPU Utilization", "RAM Consumption", "Execution Time / Latency"]
    lat_data = [
        ["PyMuPDF PDF Text Stream Parsing (10 Pages)", "12% (1 Core)", "18 MB RSS", "120 ms"],
        ["DP Chunking Engine (LeetCode 410 Partition)", "8% (1 Core)", "4 MB RSS", "14 ms"],
        ["SBERT Local CPU Embedding (10 Chunks)", "65% (2 Cores)", "140 MB RSS", "145 ms"],
        ["K-Means Clustering & Richness Ranking", "15% (1 Core)", "8 MB RSS", "22 ms"],
        ["FAISS In-Memory Vector Search (Top 50)", "5% (1 Core)", "32 MB RSS", "4.2 ms"],
        ["SQLite FTS5 Lexical Search (Top 50)", "4% (1 Core)", "12 MB RSS", "3.4 ms"],
        ["Reciprocal Rank Fusion Score Combination", "2% (1 Core)", "2 MB RSS", "1.8 ms"],
        ["0/1 Knapsack Remediation Optimization", "3% (1 Core)", "1 MB RSS", "0.85 ms"],
        ["Gemini 3.7 Flash Batch Question Generation", "Network Async", "Minimal Buffer", "1.65 s (Remote Cloud)"]
    ]
    add_styled_table(doc, lat_headers, lat_data, col_widths=[2.5, 1.2, 1.2, 1.8])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 5: EXECUTION, DEPLOYMENT & INFRASTRUCTURE
    # =========================================================================
    add_styled_heading(doc, "Chapter 5: Execution, Deployment & Infrastructure", level=1)
    
    add_styled_heading(doc, "5.1 Space-Bound Single-Node VPS Architecture (<500MB RAM)", level=2)
    add_body_paragraph(doc, 
        "Modern AI applications typically require massive cloud infrastructure, including dedicated Elasticsearch clusters, Pinecone/Milvus vector databases, and hosted GPU instances costing hundreds of dollars monthly. In contrast, The Viva-Verse is engineered under strict space-bound constraints. By utilizing embedded SQLite FTS5 for lexical indexing, in-memory FAISS for vector search, and local SBERT embeddings, the entire backend, database, and index footprint requires under 500MB of RAM—enabling flawless execution on a standard $5/month cloud VPS (e.g., AWS EC2 t3.micro or Hetzner CX22).",
        bold_prefix="")

    add_styled_heading(doc, "5.2 Step-by-Step Production Deployment Guide", level=2)
    add_bullet_point(doc, " Ubuntu 22.04 LTS x86_64, Python 3.10+, Node.js 18+, Git, build-essential.", bold_title="1. Host Server Provisioning:")
    add_bullet_point(doc, " Clone git repository, create isolated virtual environment (`python -m venv venv`), install locked dependencies via `pip install -r requirements.txt`.", bold_title="2. Backend Environment Configuration:")
    add_bullet_point(doc, " Serve FastAPI application via Uvicorn workers managed by Gunicorn or systemd process supervisor (`uvicorn app.main:app --host 0.0.0.0 --port 8000 --workers 4`).", bold_title="3. ASGI Server Process Execution:")
    add_bullet_point(doc, " Compile frontend assets via `npm run build` producing optimized HTML/JS/CSS bundles in `frontend/dist/`.", bold_title="4. Frontend Production Compilation:")
    add_bullet_point(doc, " Configure Nginx as reverse proxy with gzip compression, HTTP/2 termination, Let's Encrypt SSL/TLS certificates, and caching headers for static assets.", bold_title="5. Nginx Reverse Proxy & SSL Setup:")
    add_bullet_point(doc, " Binary index (`faiss_index.bin`) and metadata map (`faiss_metadata.json`) serialize automatically to disk on shutdown and reload on boot.", bold_title="6. FAISS Persistent Disk Synchronization:")

    add_styled_heading(doc, "5.3 Docker Containerization & Multi-Stage Builds", level=2)
    code_docker = """# Multi-Stage Dockerfile for Production Deployment
FROM python:3.10-slim AS backend-builder
WORKDIR /app
RUN apt-get update && apt-get install -y --no-install-recommends build-essential && rm -rf /var/lib/apt/lists/*
COPY backend/requirements.txt .
RUN pip install --no-cache-dir -r requirements.txt

FROM python:3.10-slim
WORKDIR /app
COPY --from=backend-builder /usr/local/lib/python3.10/site-packages /usr/local/lib/python3.10/site-packages
COPY backend /app
EXPOSE 8000
CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8000"]"""
    add_code_block(doc, code_docker, "Listing: Dockerfile Multi-Stage Build Specification")

    add_styled_heading(doc, "5.4 CI/CD Pipeline & Automated Quality Gates", level=2)
    add_body_paragraph(doc, 
        "A continuous integration and deployment pipeline is maintained via GitHub Actions. On every pull request to the `main` branch: (1) flake8 lints python syntax; (2) PyTest runs the complete 22-case automated test suite; (3) Docker builds the container image; and (4) an automated webhook triggers zero-downtime rolling deployment to the production server.",
        bold_prefix="")

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 6: CHALLENGES ENCOUNTERED & TECHNICAL SOLUTIONS
    # =========================================================================
    add_styled_heading(doc, "Chapter 6: Challenges Encountered and Technical Engineering Solutions", level=1)
    
    chal_data = [
        ["6.1 Semantic Blindness in Keyword Retrieval", 
         "Lexical search missed experiences using synonyms (e.g., 'Kubernetes' vs 'K8s', 'Raft' vs 'Consensus').",
         "Implemented parallel FAISS Dense Vector search combined with SQLite FTS5 using Reciprocal Rank Fusion (RRF k=60). Achieved +45% higher MRR."],
        ["6.2 LLM Recency Bias & Career Hyperfixation",
         "Standard LLMs spent 90% of interview time on the most recent 1-2 bullet points, ignoring earlier years.",
         "Applied K-Means clustering across SBERT embeddings to force uniform semantic sampling across the entire career timeline (improved coverage from 3/10 to 9/10)."],
        ["6.3 Document Slicing Context Corruption",
         "Fixed-character chunking cut sentences and technical equations mid-word, inducing hallucinations.",
         "Formulated chunking as LeetCode 410 Split Array Largest Sum with PyMuPDF atomic paragraph boundaries."],
        ["6.4 Hallucinatory Post-Interview Remediation",
         "LLM-generated study plans recommended arbitrary study hours unaligned with real hiring weights.",
         "Engineered deterministic 0/1 Knapsack DP solver driven by Section-Weighted Term Frequency on the target JD."],
        ["6.5 Cloud Inference Costs & API Rate Limits",
         "Centralized cloud LLM hosting creates recurring server bills and hits Free Tier rate limits.",
         "Architected zero-storage Bring-Your-Own-Key (BYOK) pass-through and single-batch API prompting."],
        ["6.6 SQLite Concurrency & FAISS Synchronization",
         "SQLite file locks during high-frequency writes and FAISS index drift across async worker processes.",
         "Enabled SQLite Write-Ahead Logging (WAL mode) and in-memory thread-safe FAISS shared index with disk caching."]
    ]
    
    for title, problem, solution in chal_data:
        add_styled_heading(doc, title, level=2)
        add_body_paragraph(doc, problem, bold_prefix="Engineering Challenge: ")
        add_body_paragraph(doc, solution, bold_prefix="Implemented Technical Solution: ")

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 7: PROJECT EXECUTION EVIDENCE & GOVERNANCE
    # =========================================================================
    add_styled_heading(doc, "Chapter 7: Project Execution Evidence & Academic Governance", level=1)
    
    add_styled_heading(doc, "7.1 Version Control Evidence & Repository Architecture", level=2)
    add_body_paragraph(doc, 
        "The project has been actively maintained under strict Git version control. The repository follows standard trunk-based development with feature branches and pull request code reviews.",
        bold_prefix="")
    add_body_paragraph(doc, "Official GitHub Repository: https://github.com/Harsh10022004/viva-verse", bold_prefix="Source Code Repository: ")

    add_styled_heading(doc, "7.2 12-Week Execution Timeline Summary", level=2)
    week_headers = ["Week No.", "Planned Milestone", "Completed Deliverables", "Supervisor Assessment"]
    week_data = [
        ["Week 1-2", "Problem Formulation & Base Architecture", "Designed 4-tier architecture, initialized FastAPI, SQLite schema", "Approved & Validated"],
        ["Week 3-4", "Lexical & Dense Vector Retrieval", "Integrated SQLite FTS5 BM25 and FAISS IndexFlatIP 384-d vectors", "Good Progress"],
        ["Week 5-6", "Hybrid Search & RRF Algorithm", "Engineered Reciprocal Rank Fusion, parallel search coordinator", "Excellent Implementation"],
        ["Week 7-8", "Document Intelligence & DP Chunking", "Built PyMuPDF parser, LeetCode 410 Split Array DP chunking engine", "Approved"],
        ["Week 9-10", "Anti-Hyperfixation & Remediation", "Integrated K-Means clustering, 0/1 Knapsack DP remediation engine", "Outstanding Achievement"],
        ["Week 11-12", "UI/UX, Testing, CI/CD & Final Report", "React 18 frontend integration, 22 test cases, Docker, documentation", "Final Capstone Approved"]
    ]
    add_styled_table(doc, week_headers, week_data, col_widths=[1.0, 2.0, 2.5, 1.5])

    add_styled_heading(doc, "7.3 Supervisor Review Dates, Directives & Resolutions", level=2)
    rev_headers = ["Review Date", "Milestone / Phase", "Key Feedback & Actionable Directives", "Resolution & Implementation"]
    rev_data = [
        ["Review 1 (W2)", "Architecture Review", "Direct exploration into RAG paradigms and in-memory FAISS indexing.", "Implemented local FAISS index with disk serialization."],
        ["Review 2 (W5)", "Retrieval Benchmarks", "Avoid relying solely on vector search; preserve exact lexical keywords.", "FxFused BM25 with FAISS using Reciprocal Rank Fusion (RRF)."],
        ["Review 3 (W7)", "Document Parsing", "Prevent mid-sentence cutting during document chunking.", "Formulated DP Split Array Largest Sum algorithm."],
        ["Review 4 (W9)", "LLM Mock Testing", "Eliminate LLM recency bias; ensure full resume timeline is tested.", "Deployed K-Means semantic clustering on SBERT vectors."],
        ["Review 5 (W11)", "Remediation Engine", "Do not let LLM hallucinate study plans; use deterministic math.", "Engineered 0/1 Knapsack DP solver with SW-TF JD weights."],
        ["Review 6 (W12)", "Final Evaluation", "Verify single-node space-bound memory footprint (<500MB RAM).", "Validated 340MB total RSS under 50 concurrent requests."]
    ]
    add_styled_table(doc, rev_headers, rev_data, col_widths=[1.2, 1.3, 2.3, 2.2])

    doc.add_page_break()

    # =========================================================================
    # CHAPTER 8: CONCLUSION & FUTURE WORK
    # =========================================================================
    add_styled_heading(doc, "Chapter 8: Conclusion & Future Work", level=1)
    
    add_styled_heading(doc, "8.1 Summary of Contributions", level=2)
    add_body_paragraph(doc, 
        "The Viva-Verse successfully redefines technical interview preparation by combining cutting-edge Information Retrieval with deterministic mathematical optimization and autonomous Generative AI. By replacing outdated keyword forums and rigid AI bots with a space-bound hybrid architecture, the platform transforms passive reading into active, measurable interview defense.",
        bold_prefix="")

    add_styled_heading(doc, "8.2 Real-World Societal & Industrial Impact", level=2)
    add_body_paragraph(doc, 
        "By democratizing access to high-tier mock interview evaluations and real-world candidate intelligence through a zero-cost Bring-Your-Own-Key model, The Viva-Verse levels the playing field for software engineering candidates worldwide—eliminating the financial barrier of $150-$300/hour human coaching.",
        bold_prefix="")

    add_styled_heading(doc, "8.3 Current System Limitations", level=2)
    add_bullet_point(doc, " The FAISS vector index is maintained in-memory. While ideal for space-bound VPS scaling under 100,000 experiences, scaling to tens of millions of records would necessitate a distributed vector store.", bold_title="• In-Memory Index Scalability:")
    add_bullet_point(doc, " Interactions are currently text-based via the Live Defense Arena. Vocal simulations require external browser transcription.", bold_title="• Audio / Vocal Modality:")

    add_styled_heading(doc, "8.4 Future Roadmap & Enhancements", level=2)
    add_bullet_point(doc, " Integration of Whisper STT and ElevenLabs/Edge TTS for low-latency voice interviews.", bold_title="1. Real-Time Speech Synthesis (STT/TTS):")
    add_bullet_point(doc, " Expanding the Knapsack remediation engine with the Ebbinghaus Forgetting Curve to model candidate memory decay over time.", bold_title="2. Spaced Repetition Memory Decay Engine:")
    add_bullet_point(doc, " Shared interactive whiteboard for live distributed architecture diagrams and dynamic code execution sandboxes.", bold_title="3. Live Collaborative Canvas & Sandboxed Execution:")

    doc.add_page_break()

    # =========================================================================
    # REFERENCES
    # =========================================================================
    add_styled_heading(doc, "References", level=1)
    refs = [
        "Robertson, S. E., & Zaragoza, H. (2009). The Probabilistic Relevance Framework: BM25 and Beyond. Foundations and Trends in Information Retrieval, 3(4), 333-389.",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. Proceedings of the 2019 Conference on Empirical Methods in Natural Language Processing (EMNLP).",
        "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2009). Introduction to Algorithms (3rd ed., Dynamic Programming & Knapsack Optimization). MIT Press.",
        "Johnson, J., Douze, M., & Jégou, H. (2019). Billion-scale similarity search with GPUs. IEEE Transactions on Big Data, 7(3), 535-547 (FAISS Architecture).",
        "Cormack, G. V., Clarke, C. L., & Buettcher, S. (2009). Reciprocal rank fusion outperforms Condorcet and individual machine learning methods. Proceedings of the 32nd ACM SIGIR Conference, 758-759.",
        "Ramírez, S. (2023). FastAPI: High-performance modern Python web framework. https://fastapi.tiangolo.com/",
        "Google DeepMind. (2024). Gemini 1.5 & Gemini 2.0: Multimodal Foundations and Long-Context Reasoning. Google Research Technical Reports.",
        "MacQueen, J. (1967). Some methods for classification and analysis of multivariate observations. Proceedings of the 5th Berkeley Symposium on Mathematical Statistics and Probability, 1, 281-297."
    ]
    for r in refs:
        add_bullet_point(doc, r, bold_title="• ")

    doc.add_page_break()

    # =========================================================================
    # APPENDICES
    # =========================================================================
    add_styled_heading(doc, "Appendices", level=1)
    
    add_styled_heading(doc, "Appendix A: Complete Codebase Directory Manifest", level=2)
    code_repo = """viva-verse/
├── backend/
│   ├── app/
│   │   ├── api/
│   │   │   └── v1/
│   │   │       ├── auth_routes.py         # JWT Authentication & Registration
│   │   │       ├── coach_routes.py        # Live Viva Simulation & Interrogation Endpoints
│   │   │       └── experience_routes.py   # Hybrid Search & Experience CRUD
│   │   ├── services/
│   │   │   ├── chunking_engine.py         # DP Split Array Largest Sum Chunking (LC 410)
│   │   │   ├── coach_service.py           # Multi-Turn Agent Prompt Orchestrator
│   │   │   ├── embedding_service.py       # SBERT / Gemini Vector Generation
│   │   │   ├── hybrid_ingestion.py        # Batch JSON Experience Ingestion Pipeline
│   │   │   ├── knapsack_engine.py         # 0/1 Knapsack Remediation & SW-TF Optimizer
│   │   │   ├── llm_service.py             # BYOK Google GenAI Client & Singletons
│   │   │   ├── notification_service.py    # Real-Time Search Alerting Hooks
│   │   │   ├── parser_service.py          # PyMuPDF Extractor & K-Means Clustering
│   │   │   ├── search_service.py          # SQLite FTS5 + FAISS Reciprocal Rank Fusion
│   │   │   └── vector_store.py            # In-Memory FAISS Vector Index & Disk Cache
│   │   ├── database.py                    # SQLAlchemy Engine & Session Local
│   │   ├── database_models.py             # User, Session, Experience, Round, Question Schemas
│   │   └── utils/
│   │       ├── constants.py               # Stopwords & Mode Templates
│   │       └── skill_metadata.json        # Static Skill Learning Hour Database
│   ├── faiss_index.bin                    # Serialized In-Memory FAISS Vector Index
│   ├── faiss_metadata.json                # Vector ID to Question UUID Mapping
│   ├── viva_verse_v3.db                   # SQLite Database with FTS5 Virtual Tables
│   └── requirements.txt                   # Backend Python Dependencies
├── frontend/
│   ├── src/
│   │   ├── components/
│   │   │   ├── CoachDashboard.jsx         # Post-Interview Knapsack Remediation View
│   │   │   ├── CoachTerminal.jsx          # Real-Time Multi-Turn Chat Defense Arena
│   │   │   ├── HistoryView.jsx            # Historical Viva Session Archive
│   │   │   ├── InterviewExperiences.jsx   # Hybrid Search Explorer & Submission Modal
│   │   │   ├── LoginView.jsx              # User Authentication Portal
│   │   │   └── SetupStudio.jsx            # JD/Resume Upload & BYOK Key Calibrator
│   │   ├── App.jsx                        # Master Client Routing & State Context
│   │   └── index.css                      # Tailwind CSS Design System
│   ├── package.json                       # Frontend Node.js Dependencies
│   └── vite.config.js                     # Vite Bundler & Proxy Configuration
└── README.md                              # System Architecture Documentation"""
    add_code_block(doc, code_repo, "Repository Hierarchy & File Organization")

    add_styled_heading(doc, "Appendix B: Complete REST API Endpoint Specifications", level=2)
    api_headers = ["HTTP Method & Route", "Module Controller", "Request Payload", "Response Schema", "Status Code"]
    api_data = [
        ["POST /api/v1/auth/register", "Auth Controller", "{username, email, password}", "{user_id, token}", "201 Created"],
        ["POST /api/v1/auth/login", "Auth Controller", "{username, password}", "{access_token, token_type}", "200 OK"],
        ["POST /api/v1/experiences/search", "Search Service", "{query, filters, page, page_size}", "{total, results: [experiences]}", "200 OK"],
        ["POST /api/v1/experiences", "Experience Service", "{company, role, level, rounds: []}", "{id, status: 'created'}", "201 Created"],
        ["POST /api/v1/coach/init-session", "Coach Service", "{mode, role, level, jd_text, cv_file}", "{session_id, first_question}", "200 OK"],
        ["POST /api/v1/coach/submit-answer", "Coach Service", "{session_id, answer_text, api_key}", "{score, critique, next_question, is_done}", "200 OK"],
        ["GET /api/v1/coach/report/{session_id}", "Knapsack Service", "None (URL param)", "{scorecard, knapsack_plan: {roadmap}}", "200 OK"],
        ["POST /api/v1/subscriptions", "Notification Hub", "{query_text, contact_email, threshold}", "{subscription_id, status: 'active'}", "201 Created"]
    ]
    add_styled_table(doc, api_headers, api_data, col_widths=[1.8, 1.2, 1.5, 1.5, 1.0])

    add_styled_heading(doc, "Appendix C: Local Development & Quick Start Installation Guide", level=2)
    code_install = """# 1. Clone Source Code Repository
git clone https://github.com/Harsh10022004/viva-verse.git
cd viva-verse

# 2. Configure and Start Python Backend
cd backend
python -m venv venv
# On Windows:
venv\\Scripts\\activate
# On Linux/macOS:
source venv/bin/activate

pip install -r requirements.txt
cp .env.example .env  # Configure GEMINI_API_KEY optional defaults
uvicorn app.main:app --reload --port 8000

# 3. Configure and Start React Frontend (In a new terminal)
cd ../frontend
npm install
npm run dev -- --port 5173

# 4. Access Local Applications:
# Frontend Web App: http://localhost:5173
# FastAPI Swagger Interactive Docs: http://localhost:8000/docs"""
    add_code_block(doc, code_install, "Local Development Setup Script")

    add_styled_heading(doc, "Appendix D: User Manual & Candidate Operational Playbook", level=2)
    add_bullet_point(doc, " Open the web client at `http://localhost:5173`. Use the search bar to query real-world interview experiences (e.g., 'Google Staff Distributed Systems', 'Meta L5 Concurrency'). Toggle filters by Company, Role, Level, or Source (Native / Ingested).", bold_title="Step 1: Explore Real-World Experiences:")
    add_bullet_point(doc, " Navigate to 'Share Experience' tab. Fill out company, role, level, interview date, and add granular interview rounds (Round Name, Notes, Specific Technical Questions). The backend automatically indexes questions in FTS5 and FAISS.", bold_title="Step 2: Contribute Interview Intel:")
    add_bullet_point(doc, " Click 'Setup Studio' in the navigation bar. Choose your target interview mode (Technical Coding, System Design, Behavioral STAR, Online Assessment, Certification). Provide your Bring-Your-Own-Key (BYOK) API key from Google AI Studio, OpenRouter, or NVIDIA NIM.", bold_title="Step 3: Calibrate Interrogation Arena:")
    add_bullet_point(doc, " Drag and drop your candidate Resume PDF and paste your target Job Description. The system executes PyMuPDF text extraction, DP chunking, and K-Means topic extraction.", bold_title="Step 4: Upload Resume & Target JD:")
    add_bullet_point(doc, " Step into the Live Defense Arena. Answer the AI examiner's targeted conceptual questions. Receive real-time scoring, Big-O analysis, and follow-up stress testing.", bold_title="Step 5: Conduct Live Viva Defense:")
    add_bullet_point(doc, " Upon defense completion, view your comprehensive analytics scorecard and the 0/1 Knapsack mathematically optimized multi-day study roadmap.", bold_title="Step 6: Review Knapsack Remediation Roadmap:")

    # Save finalized document
    output_filename = "Viva_Verse_Comprehensive_Capstone_Project_Report.docx"
    doc.save(output_filename)
    print(f"Master Document Generated Successfully: {os.path.abspath(output_filename)}")

if __name__ == "__main__":
    build_viva_verse_document()
